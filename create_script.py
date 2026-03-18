from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create Script Document
doc = Document()

# Title
title = doc.add_heading('Paddy Power IMC Presentation - Group Script', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('10-Minute Presentation | 6 Speakers')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(12)
subtitle_format.font.italic = True

doc.add_paragraph()

# Introduction
doc.add_heading('SPEAKER 1: Introduction & Brand Overview (1.5 minutes)', 1)

script1 = """
[0:00-1:30]

"Good morning/afternoon everyone. Today, our group will be analysing the Integrated Marketing Communications strategy of Paddy Power, one of Ireland's most recognisable betting brands.

Paddy Power was established in 1988 through the merger of three Irish bookmakers. Since then, it has built its reputation on a distinctive IMC strategy centred on what they call 'mischief' - a unified, irreverent brand personality that is consistently applied across all channels.

The key element of Paddy Power's IMC strategy is maintaining a consistent, bold, and humorous brand message. Their cheeky 'lad culture' persona positions them as an entertainment provider rather than just a traditional bookmaker. By 2026, this strategy has evolved to balance high-energy entertainment with a strong focus on responsible gambling and technological integration.

Our presentation will examine how Paddy Power uses six key IMC tools: Advertising, Public Relations, Social Media, Direct Marketing, Sales Promotion, and Sponsorship - and how these work together to create a cohesive brand experience."
"""
doc.add_paragraph(script1)

# Advertising
doc.add_heading('SPEAKER 2: Advertising (1.5 minutes)', 1)

script2 = """
[1:30-3:00]

"Paddy Power utilises a mix of traditional and digital advertising channels to reach their audience. Their television commercials feature high-profile celebrities and humorous scenarios that reinforce the brand's cheeky personality.

The standout example is their 'Come Out and Play' campaign, launched in October 2025. Created by the agency BBH, this campaign stars Danny Dyer, Coleen Rooney, Gemma Collins, and Peter Crouch. The concept is described as 'Queen Vic meets Las Vegas' - a dreamlike casino scenario where Danny Dyer acts as 'The House,' a cockney casino guide leading viewers through a star-studded tour.

The campaign launched across TV, Video on Demand, and social channels, using self-deprecating humour and Dyer's trademark patter to create a memorable, enjoyable experience. This demonstrates how Paddy Power uses high-impact advertising with celebrity endorsements to enhance brand recall across multiple platforms."
"""
doc.add_paragraph(script2)

# PR and Social Media
doc.add_heading('SPEAKER 3: Public Relations & Social Media (2 minutes)', 1)

script3 = """
[3:00-5:00]

"Paddy Power frequently creates publicity stunts designed to attract media attention and generate earned media coverage. Let me share two key examples.

First, the 'Hollywood Sign' stunt from the 2010 Ryder Cup at Celtic Manor Resort. Paddy Power created a giant hillside sign reading 'Paddy Power' instead of 'Hollywood.' This generated widespread news coverage and social media sharing, with millions seeing the brand name without traditional advertising costs.

Second, the 'Even Bigger 180' campaign from 2024 to 2026, linked to the PDC World Darts Championship. Paddy Power promised to donate £1,000 for every '180' score to Prostate Cancer UK. The campaign combined charitable donations with QR codes in venues and national newspaper takeovers, raising over £1 million and encouraging thousands of men to check their cancer risk.

On social media, Paddy Power uses platforms like Instagram, X, and TikTok to react to live sports, share humorous content and memes, and engage with fans. For example, before a major Katie Taylor boxing fight, they posted a humorous video warning fans about toxic comments. This real-time engagement keeps the brand relevant and encourages sharing."
"""
doc.add_paragraph(script3)

# Direct Marketing and Sales Promotion
doc.add_heading('SPEAKER 4: Direct Marketing & Sales Promotion (2 minutes)', 1)

script4 = """
[5:00-7:00]

"Paddy Power communicates directly with customers through email, SMS, and push notifications in their app. These channels deliver personalised betting offers and updates, helping maintain customer relationships and encourage repeat usage.

Their sales promotion strategy includes the 'Money Back Special' and 'Justice Payouts' - refunding bets deemed unfair. These are integrated across all platforms to build customer trust and reinforce a 'customer-friendly' reputation.

The flagship promotion is Paddy's Rewards Club. Members who place 5 or more bets of £5 or €5 at odds of 1/2 or higher between Monday and Sunday receive rewards the following Monday. These include free bets ranging from £5 to £50, or Power Up tokens that boost odds on eligible bets.

Importantly, this applies across all channels - online, mobile, phone, text, and even shop bets with a linked Play Card. In 2025, the program was updated to require opt-in, ensuring engaged participation. New customers can also access offers like betting £5 and receiving £30 in free bets for the jumps season."
"""
doc.add_paragraph(script4)

# Sponsorship and Integration
doc.add_heading('SPEAKER 5: Sponsorship & Integrated Approach (1.5 minutes)', 1)

script5 = """
[7:00-8:30]

"Paddy Power's sponsorship strategy demonstrates effective IMC integration. Their sponsorship of the PDC World Darts Championship, combined with the 'Even Bigger 180' campaign, shows how sponsorship can be linked with charitable giving, QR code activation, and media coverage to maximise impact.

The brand also sponsors First Dates, extending into 2026, to build brand warmth and cultural relevance beyond sports betting.

All these communication tools work together to deliver a single, unified brand message. A campaign might start with a controversial stunt, gain attention through PR coverage, be amplified on social media, and then lead customers to targeted promotions through email or the mobile app.

This creates what marketing professionals call the 'pinball effect' - where consumers enter at unpredictable points, requiring concise, contextual messaging that drives emotional connections and conversions. Whether customers encounter Paddy Power through a billboard, a TV ad, a tweet, or an email, they receive consistent brand messaging."
"""
doc.add_paragraph(script5)

# Conclusion
doc.add_heading('SPEAKER 6: Responsible Gambling & Conclusion (1.5 minutes)', 1)

script6 = """
[8:30-10:00]

"A significant evolution in Paddy Power's 2026 strategy is the integration of responsible gambling tools directly into their marketing messaging. This includes AI-driven time alerts and self-exclusion options, aligning with stricter global regulations while demonstrating corporate responsibility.

The brand has also focused on 'transcending' the brick-and-mortar experience by integrating digital verification and AI-driven personalisation into retail spaces, creating a seamless omni-channel experience.

In conclusion, Paddy Power's IMC strategy demonstrates how a consistent brand personality - their 'mischief' positioning - executed across multiple channels with integrated messaging, can create a distinctive market position. By balancing entertainment value with responsible practices, the brand maintains relevance while adapting to evolving regulatory and consumer expectations.

The key lesson is that successful IMC requires not just using multiple channels, but ensuring they work together to deliver a unified brand experience - something Paddy Power achieves through their cheeky, provocative, yet increasingly responsible approach.

Thank you for listening. We'd be happy to take any questions."
"""
doc.add_paragraph(script6)

# Timing notes
doc.add_page_break()
doc.add_heading('Timing Summary', 1)

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Speaker'
hdr_cells[1].text = 'Topic'
hdr_cells[2].text = 'Duration'

# Content
row1 = table.rows[1].cells
row1[0].text = 'Speaker 1'
row1[1].text = 'Introduction & Brand Overview'
row1[2].text = '1.5 min (0:00-1:30)'

row2 = table.rows[2].cells
row2[0].text = 'Speaker 2'
row2[1].text = 'Advertising'
row2[2].text = '1.5 min (1:30-3:00)'

row3 = table.rows[3].cells
row3[0].text = 'Speaker 3'
row3[1].text = 'PR & Social Media'
row3[2].text = '2 min (3:00-5:00)'

row4 = table.rows[4].cells
row4[0].text = 'Speaker 4'
row4[1].text = 'Direct Marketing & Sales Promotion'
row4[2].text = '2 min (5:00-7:00)'

row5 = table.rows[5].cells
row5[0].text = 'Speaker 5'
row5[1].text = 'Sponsorship & Integration'
row5[2].text = '1.5 min (7:00-8:30)'

row6 = table.rows[6].cells
row6[0].text = 'Speaker 6'
row6[1].text = 'Responsible Gambling & Conclusion'
row6[2].text = '1.5 min (8:30-10:00)'

# Notes
doc.add_paragraph()
doc.add_heading('Presentation Notes:', 2)
doc.add_paragraph('• Total duration: 10 minutes', style='List Bullet')
doc.add_paragraph('• Each speaker should practice their section to ensure smooth transitions', style='List Bullet')
doc.add_paragraph('• Maintain eye contact with camera for recording', style='List Bullet')
doc.add_paragraph('• Use the PowerPoint slides as visual support', style='List Bullet')
doc.add_paragraph('• Speak clearly and at a moderate pace', style='List Bullet')

# Save document
doc.save('C:/Users/Karen/.openclaw/workspace/Paddy_Power_IMC_Script.docx')
print('Script document created: Paddy_Power_IMC_Script.docx')